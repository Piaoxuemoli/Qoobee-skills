"""Export Markdown paper to DOCX with IEEE single-column formatting.

Font/style specs:
  - Paper: A4, margins 2.54cm top/bottom, 3.17cm left/right
  - Title: SimHei 16pt, centered, bold
  - Author/affiliation: SimSun 12pt, centered
  - Abstract: SimSun 10.5pt, italic, indented 1cm
  - Keywords: SimSun 10.5pt
  - Body: SimSun 12pt (小四), 1.5x line spacing
  - Heading 1: SimHei 14pt, bold
  - Heading 2: SimHei 12pt, bold
  - References: SimSun 10.5pt (五号)
  - Header: paper title (SimSun 9pt)
  - Footer: page number centered

Usage:
    python export_docx.py \
        --input "outputs/paper/04_final/final_paper.md" \
        --output "outputs/paper/05_exports/final_paper.docx"
"""
from __future__ import annotations
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: requires python-docx. pip install python-docx", file=sys.stderr)
    sys.exit(1)


def _set_font(run, name_cn="SimSun", name_en="Times New Roman",
              size=Pt(12), bold=False, italic=False):
    """Set font properties on a run."""
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name_en
    # set East Asian font
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)


def _set_paragraph_spacing(paragraph, before=Pt(0), after=Pt(0),
                           line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line_spacing


def _add_title(doc, text):
    """Add paper title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_font(run, "SimHei", "Arial", Pt(16), bold=True)
    _set_paragraph_spacing(p, before=Pt(24), after=Pt(12))


def _add_author(doc, text):
    """Add author line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_font(run, "SimSun", "Times New Roman", Pt(12))
    _set_paragraph_spacing(p, after=Pt(6))


def _add_abstract(doc, text):
    """Add abstract block."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.right_indent = Cm(1)
    run_label = p.add_run("摘要：")
    _set_font(run_label, "SimHei", "Arial", Pt(10.5), bold=True)
    run_text = p.add_run(text)
    _set_font(run_text, "SimSun", "Times New Roman", Pt(10.5), italic=True)
    _set_paragraph_spacing(p, before=Pt(12), after=Pt(6), line_spacing=1.25)


def _add_keywords(doc, text):
    """Add keywords line."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.right_indent = Cm(1)
    run_label = p.add_run("关键词：")
    _set_font(run_label, "SimHei", "Arial", Pt(10.5), bold=True)
    run_text = p.add_run(text)
    _set_font(run_text, "SimSun", "Times New Roman", Pt(10.5))
    _set_paragraph_spacing(p, after=Pt(12), line_spacing=1.25)


def _add_heading1(doc, text):
    """Add level-1 heading (黑体 14pt bold)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, "SimHei", "Arial", Pt(14), bold=True)
    _set_paragraph_spacing(p, before=Pt(18), after=Pt(6))


def _add_heading2(doc, text):
    """Add level-2 heading (黑体 12pt bold)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, "SimHei", "Arial", Pt(12), bold=True)
    _set_paragraph_spacing(p, before=Pt(12), after=Pt(6))


def _add_body(doc, text):
    """Add body text paragraph (宋体 12pt, 1.5x line spacing)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, "SimSun", "Times New Roman", Pt(12))
    _set_paragraph_spacing(p, after=Pt(6), line_spacing=1.5)
    p.paragraph_format.first_line_indent = Cm(0.74)  # 2 char indent


def _add_reference(doc, text):
    """Add reference entry (宋体 10.5pt)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, "SimSun", "Times New Roman", Pt(10.5))
    _set_paragraph_spacing(p, after=Pt(2), line_spacing=1.25)


def _setup_page(doc):
    """Set A4 page with IEEE-style margins."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def _add_header(doc, title):
    """Add page header with paper title."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_font(run, "SimSun", "Times New Roman", Pt(9))


def _add_footer_page_number(doc):
    """Add centered page number in footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # add page number field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def _parse_markdown(md_text: str) -> list:
    """Parse markdown into structured blocks.

    Returns list of (type, content) tuples:
      ("title", text)       -- # heading
      ("h1", text)          -- ## heading
      ("h2", text)          -- ### heading
      ("abstract", text)    -- **摘要：** content
      ("keywords", text)    -- **关键词：** content
      ("body", text)        -- regular paragraph
      ("reference", text)   -- [N] citation
      ("blank", "")         -- empty line
    """
    blocks = []
    lines = md_text.splitlines()
    in_references = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            blocks.append(("blank", ""))
            continue

        # title: # Title
        if stripped.startswith("# ") and not stripped.startswith("## "):
            blocks.append(("title", stripped[2:].strip()))
            continue

        # h1: ## Heading
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            if "参考文献" in text or "references" in text.lower():
                in_references = True
            blocks.append(("h1", text))
            continue

        # h2: ### Heading
        if stripped.startswith("### "):
            blocks.append(("h2", stripped[4:].strip()))
            continue

        # abstract
        if stripped.startswith("**摘要") or stripped.startswith("**Abstract"):
            text = re.sub(r"^\*\*[^*]+\*\*\s*", "", stripped)
            blocks.append(("abstract", text))
            continue

        # keywords
        if stripped.startswith("**关键词") or stripped.startswith("**Keywords"):
            text = re.sub(r"^\*\*[^*]+\*\*\s*", "", stripped)
            blocks.append(("keywords", text))
            continue

        # reference line: [N] ...
        if re.match(r"^\[\d+\]", stripped):
            blocks.append(("reference", stripped))
            continue

        # skip markdown bold-only lines (already handled)
        if stripped.startswith("**") and stripped.endswith("**"):
            continue

        # regular body text
        blocks.append(("body", stripped))

    return blocks


def export_docx(md_path: str, docx_path: str) -> str:
    """Convert Markdown paper to DOCX with IEEE single-column formatting."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    blocks = _parse_markdown(md_text)

    doc = Document()
    _setup_page(doc)

    # extract title for header
    title_text = ""
    for btype, btext in blocks:
        if btype == "title":
            title_text = btext
            break

    if title_text:
        _add_header(doc, title_text)
    _add_footer_page_number(doc)

    # remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    # render blocks
    for btype, btext in blocks:
        if btype == "blank":
            continue
        elif btype == "title":
            _add_title(doc, btext)
        elif btype == "h1":
            _add_heading1(doc, btext)
        elif btype == "h2":
            _add_heading2(doc, btext)
        elif btype == "abstract":
            _add_abstract(doc, btext)
        elif btype == "keywords":
            _add_keywords(doc, btext)
        elif btype == "reference":
            _add_reference(doc, btext)
        elif btype == "body":
            _add_body(doc, btext)

    # save
    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    return docx_path


def main():
    parser = argparse.ArgumentParser(description="Export Markdown to IEEE-style DOCX")
    parser.add_argument("--input", required=True, help="Input .md file")
    parser.add_argument("--output", required=True, help="Output .docx file")
    args = parser.parse_args()

    result = export_docx(args.input, args.output)
    print(f"DOCX exported: {result}")


if __name__ == "__main__":
    main()
